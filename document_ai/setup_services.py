#!/usr/bin/env python3
"""
Script to populate all service files for Document AI
"""
import os

def create_file(path, content):
    """Create a file with the given content"""
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f'Created: {path}')

def main():
    base = os.path.dirname(os.path.abspath(__file__))
    
    # PDF Service
    create_file(f'{base}/app/services/pdf_service.py', '''import os
from typing import List, Dict, Any, Optional
from pathlib import Path
from dataclasses import dataclass

try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from unstructured.partition.pdf import partition_pdf
    HAS_UNSTRUCTURED = True
except ImportError:
    HAS_UNSTRUCTURED = False

from app.core.logger import logger


@dataclass
class PDFPage:
    page_number: int
    text: str
    elements: List[Dict[str, Any]]


class PDFService:
    
    def __init__(self):
        self.supports_unstructured = HAS_UNSTRUCTURED
        self.supports_pypdf = HAS_PYPDF
    
    def extract_text(self, file_path: str, include_layout: bool = True) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"PDF file not found: {file_path}")
        
        if include_layout and self.supports_unstructured:
            return self._extract_with_unstructured(file_path)
        elif self.supports_pypdf:
            return self._extract_with_pypdf(file_path)
        else:
            raise ImportError("No suitable PDF library available. Install pypdf or unstructured.")
    
    def _extract_with_unstructured(self, file_path: str) -> Dict[str, Any]:
        try:
            elements = partition_pdf(
                filename=file_path,
                strategy="hi_res",
                infer_table_structure=True,
                include_page_breaks=True
            )
            
            pages = []
            current_page = []
            current_page_num = 1
            
            for element in elements:
                element_dict = {
                    'type': type(element).__name__,
                    'text': str(element),
                    'metadata': element.metadata.to_dict() if hasattr(element, 'metadata') else {}
                }
                
                page_num = element_dict['metadata'].get('page_number', 1)
                
                if page_num != current_page_num and current_page:
                    pages.append({
                        'page_number': current_page_num,
                        'elements': current_page,
                        'text': ' '.join([e['text'] for e in current_page])
                    })
                    current_page = []
                    current_page_num = page_num
                
                current_page.append(element_dict)
            
            if current_page:
                pages.append({
                    'page_number': current_page_num,
                    'elements': current_page,
                    'text': ' '.join([e['text'] for e in current_page])
                })
            
            return {
                'pages': pages,
                'total_pages': len(pages),
                'full_text': ' '.join([p['text'] for p in pages]),
                'has_layout': True
            }
        except Exception as e:
            logger.error(f"Error extracting with unstructured: {e}")
            return self._extract_with_pypdf(file_path)
    
    def _extract_with_pypdf(self, file_path: str) -> Dict[str, Any]:
        reader = PdfReader(file_path)
        pages = []
        
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            pages.append({
                'page_number': i,
                'text': text,
                'elements': [{'type': 'Text', 'text': text, 'metadata': {}}],
            })
        
        return {
            'pages': pages,
            'total_pages': len(pages),
            'full_text': '\\n\\n'.join([p['text'] for p in pages]),
            'has_layout': False
        }
    
    def get_page_count(self, file_path: str) -> int:
        if self.supports_pypdf:
            reader = PdfReader(file_path)
            return len(reader.pages)
        return 0
''')
    
    # DOCX Service
    create_file(f'{base}/app/services/docx_service.py', '''import os
from typing import Dict, Any, List

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from unstructured.partition.docx import partition_docx
    HAS_UNSTRUCTURED = True
except ImportError:
    HAS_UNSTRUCTURED = False


class DOCXService:
    
    def __init__(self):
        self.supports_unstructured = HAS_UNSTRUCTURED
        self.supports_docx = HAS_DOCX
    
    def extract_text(self, file_path: str, include_layout: bool = True) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        if include_layout and self.supports_unstructured:
            return self._extract_with_unstructured(file_path)
        elif self.supports_docx:
            return self._extract_with_docx(file_path)
        else:
            raise ImportError("No suitable DOCX library available. Install python-docx or unstructured.")
    
    def _extract_with_unstructured(self, file_path: str) -> Dict[str, Any]:
        try:
            elements = partition_docx(filename=file_path)
            
            structured_elements = []
            for element in elements:
                structured_elements.append({
                    'type': type(element).__name__,
                    'text': str(element),
                    'metadata': element.metadata.to_dict() if hasattr(element, 'metadata') else {}
                })
            
            full_text = '\\n\\n'.join([e['text'] for e in structured_elements])
            
            return {
                'pages': [{
                    'page_number': 1,
                    'elements': structured_elements,
                    'text': full_text
                }],
                'total_pages': 1,
                'full_text': full_text,
                'has_layout': True
            }
        except Exception as e:
            return self._extract_with_docx(file_path)
    
    def _extract_with_docx(self, file_path: str) -> Dict[str, Any]:
        doc = Document(file_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    'type': 'Paragraph',
                    'text': para.text,
                    'metadata': {}
                })
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append({
                            'type': 'TableCell',
                            'text': cell.text,
                            'metadata': {}
                        })
        
        full_text = '\\n\\n'.join([p['text'] for p in paragraphs])
        
        return {
            'pages': [{
                'page_number': 1,
                'elements': paragraphs,
                'text': full_text
            }],
            'total_pages': 1,
            'full_text': full_text,
            'has_layout': False
        }
''')
    
    # Image Service
    create_file(f'{base}/app/services/image_service.py', '''import os
from typing import Dict, Any
from PIL import Image


class ImageService:
    
    def __init__(self):
        pass
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Image file not found: {file_path}")
        
        try:
            with Image.open(file_path) as img:
                return {
                    'width': img.width,
                    'height': img.height,
                    'format': img.format,
                    'mode': img.mode,
                    'has_transparency': img.mode in ('RGBA', 'LA', 'P')
                }
        except Exception as e:
            return {
                'error': str(e)
            }
    
    def is_likely_scanned(self, file_path: str) -> bool:
        try:
            with Image.open(file_path) as img:
                if img.mode == 'RGB':
                    pixels = list(img.getdata())
                    unique_colors = len(set(pixels[:1000]))
                    return unique_colors < 50
        except Exception:
            pass
        return False
''')
    
    # OCR Service
    create_file(f'{base}/app/services/ocr_service.py', '''import os
import logging
from typing import Dict, Any, Optional, List
from PIL import Image
import io

logger = logging.getLogger(__name__)


class OCRService:
    
    def __init__(self, provider: str = 'tesseract'):
        self.provider = provider
        self._tesseract = None
        self._easyocr = None
        self._google_vision = None
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        if self.provider == 'tesseract':
            return self._extract_with_tesseract(file_path)
        elif self.provider == 'easyocr':
            return self._extract_with_easyocr(file_path)
        elif self.provider == 'google_vision':
            return self._extract_with_google_vision(file_path)
        else:
            raise ValueError(f"Unknown OCR provider: {self.provider}")
    
    def _extract_with_tesseract(self, file_path: str) -> Dict[str, Any]:
        try:
            import pytesseract
            from pdf2image import convert_from_path
            
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                images = convert_from_path(file_path, dpi=300)
                all_text = []
                all_elements = []
                
                for i, image in enumerate(images, 1):
                    text = pytesseract.image_to_string(image)
                    all_text.append(text)
                    all_elements.append({
                        'page_number': i,
                        'text': text,
                        'type': 'OCRText'
                    })
                
                full_text = '\\n\\n'.join(all_text)
            else:
                image = Image.open(file_path)
                full_text = pytesseract.image_to_string(image)
                all_elements = [{
                    'page_number': 1,
                    'text': full_text,
                    'type': 'OCRText'
                }]
            
            return {
                'full_text': full_text,
                'elements': all_elements,
                'pages': len(all_elements),
                'provider': 'tesseract',
                'success': True
            }
        except ImportError as e:
            logger.error(f"Tesseract not available: {e}")
            return self._fallback_basic_extraction(file_path)
        except Exception as e:
            logger.error(f"Error in Tesseract OCR: {e}")
            return self._fallback_basic_extraction(file_path)
    
    def _extract_with_easyocr(self, file_path: str) -> Dict[str, Any]:
        try:
            import easyocr
            
            if self._easyocr is None:
                self._easyocr = easyocr.Reader(['en'], gpu=False)
            
            results = self._easyocr.readtext(file_path)
            
            full_text_parts = []
            elements = []
            
            for i, (bbox, text, confidence) in enumerate(results, 1):
                full_text_parts.append(text)
                elements.append({
                    'page_number': 1,
                    'text': text,
                    'confidence': confidence,
                    'type': 'OCRText'
                })
            
            full_text = ' '.join(full_text_parts)
            
            return {
                'full_text': full_text,
                'elements': elements,
                'pages': 1,
                'provider': 'easyocr',
                'success': True
            }
        except ImportError:
            logger.warning("EasyOCR not available, falling back to Tesseract")
            return self._extract_with_tesseract(file_path)
        except Exception as e:
            logger.error(f"Error in EasyOCR: {e}")
            return self._extract_with_tesseract(file_path)
    
    def _extract_with_google_vision(self, file_path: str) -> Dict[str, Any]:
        try:
            from google.cloud import vision
            from google.cloud.vision_v1 import types
            
            client = vision.ImageAnnotatorClient()
            
            with io.open(file_path, 'rb') as f:
                content = f.read()
            
            image = vision.Image(content=content)
            response = client.document_text_detection(image=image)
            
            if response.full_text_annotation:
                full_text = response.full_text_annotation.text
                pages = []
                
                for page in response.full_text_annotation.pages:
                    page_text = ''
                    for block in page.blocks:
                        for paragraph in block.paragraphs:
                            for word in paragraph.words:
                                for symbol in word.symbols:
                                    page_text += symbol.text
                                    if symbol.property.detected_break:
                                        if symbol.property.detected_break.type == 1:
                                            page_text += ' '
                                        elif symbol.property.detected_break.type == 3:
                                            page_text += '\\n'
                    pages.append({
                        'page_number': len(pages) + 1,
                        'text': page_text
                    })
                
                return {
                    'full_text': full_text,
                    'pages': pages,
                    'elements': [{'text': full_text, 'type': 'OCRText'}],
                    'provider': 'google_vision',
                    'success': True
                }
            
            return {'full_text': '', 'success': False}
        except ImportError:
            logger.warning("Google Cloud Vision not available, falling back to Tesseract")
            return self._extract_with_tesseract(file_path)
        except Exception as e:
            logger.error(f"Error in Google Vision OCR: {e}")
            return self._extract_with_tesseract(file_path)
    
    def _fallback_basic_extraction(self, file_path: str) -> Dict[str, Any]:
        try:
            image = Image.open(file_path)
            if hasattr(image, 'text'):
                metadata_text = ' '.join(image.text.values())
            else:
                metadata_text = ''
            
            return {
                'full_text': metadata_text,
                'elements': [],
                'pages': 1,
                'provider': 'fallback',
                'success': True,
                'warning': 'Limited text extraction available'
            }
        except Exception as e:
            logger.error(f"Fallback extraction failed: {e}")
            return {
                'full_text': '',
                'elements': [],
                'pages': 0,
                'provider': 'none',
                'success': False,
                'error': str(e)
            }
''')
    
    print("\\nAll service files created successfully!")

if __name__ == '__main__':
    main()
