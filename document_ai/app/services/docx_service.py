import os
from typing import Dict, Any, List

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from unstructured.partition.docx import partition_docx
    HAS_UNSTRUCTURED = True
except ImportError:
    HAS_UNSTRUCTURED = False


class DOCXService:
    
    def __init__(self):
        self.supports_unstructured = HAS_UNSTRUCTURED
        self.supports_docx = HAS_DOCX
    
    def extract_text(self, file_path: str, include_layout: bool = True) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        if include_layout and self.supports_unstructured:
            return self._extract_with_unstructured(file_path)
        elif self.supports_docx:
            return self._extract_with_docx(file_path)
        else:
            raise ImportError("No suitable DOCX library available. Install python-docx or unstructured.")
    
    def _extract_with_unstructured(self, file_path: str) -> Dict[str, Any]:
        try:
            elements = partition_docx(filename=file_path)
            
            structured_elements = []
            for element in elements:
                structured_elements.append({
                    'type': type(element).__name__,
                    'text': str(element),
                    'metadata': element.metadata.to_dict() if hasattr(element, 'metadata') else {}
                })
            
            full_text = '\n\n'.join([e['text'] for e in structured_elements])
            
            return {
                'pages': [{
                    'page_number': 1,
                    'elements': structured_elements,
                    'text': full_text
                }],
                'total_pages': 1,
                'full_text': full_text,
                'has_layout': True
            }
        except Exception as e:
            return self._extract_with_docx(file_path)
    
    def _extract_with_docx(self, file_path: str) -> Dict[str, Any]:
        doc = Document(file_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    'type': 'Paragraph',
                    'text': para.text,
                    'metadata': {}
                })
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append({
                            'type': 'TableCell',
                            'text': cell.text,
                            'metadata': {}
                        })
        
        full_text = '\n\n'.join([p['text'] for p in paragraphs])
        
        return {
            'pages': [{
                'page_number': 1,
                'elements': paragraphs,
                'text': full_text
            }],
            'total_pages': 1,
            'full_text': full_text,
            'has_layout': False
        }
